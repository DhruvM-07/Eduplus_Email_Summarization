import streamlit as st
import spacy
from prompts2 import get_prompt
from langchain_groq import ChatGroq
import os
from dotenv import load_dotenv
from langchain_core.prompts import ChatPromptTemplate
from docx import Document
from wordcloud import WordCloud
import matplotlib.pyplot as plt
import io

# Set up the Streamlit page
st.set_page_config(
    page_title="Eduplus Feedback Summarization",
    page_icon="📧",
    layout="centered",
    initial_sidebar_state="auto",
)

# Load environment variables
load_dotenv()
groq_api_key="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"

def load_llm():
    return ChatGroq(temperature=0.4, model_name="llama3-8b-8192", api_key=groq_api_key)

# Load SpaCy model
@st.cache_resource
def load_spacy_model():
    return spacy.load("en_core_web_sm")

# Function to extract emails from DOCX
def extract_emails_from_docx(file):
    document = Document(file)
    emails = []
    current_email = []
    for para in document.paragraphs:
        if para.text.strip():  # Non-empty line
            current_email.append(para.text.strip())
        elif current_email:  # Empty line marks end of email
            emails.append("\n".join(current_email))
            current_email = []
    if current_email:  # Add the last email if not added
        emails.append("\n".join(current_email))
    return emails

# Function to extract adjectives
def extract_adjectives(text):
    nlp = load_spacy_model()
    doc = nlp(text)
    return [token.text for token in doc if token.pos_ == "ADJ"]

# Function to extract CS-related sentences
def extract_cs_related_sentences(text, keywords):
    sentences = text.split(".")
    return [sentence.strip() for sentence in sentences if any(keyword in sentence.lower() for keyword in keywords)]

# Define keywords for CS-related topics
cs_keywords = [
    "algorithm", "database", "debugging", "programming", "development",
    "data science", "machine learning", "api", "backend", "frontend", "systems"
]

# Streamlit UI
st.title("Eduplus Feedback Summarization")
st.write("Upload a Word document containing internship feedback emails to analyze student performance.")

uploaded_file = st.file_uploader("Upload a DOCX file", type="docx")

llm = load_llm()

if uploaded_file:
    st.write("File uploaded successfully!")
    emails = extract_emails_from_docx(uploaded_file)
    
    if emails:
        st.write("### Extracted Emails")
        for idx, email in enumerate(emails, start=1):
            st.write(f"**Email {idx}:**")
            st.write(email)

        if st.button("Analyze Emails"):
            st.write("### Analysis Results")
            for idx, email in enumerate(emails, start=1):
                st.write(f"**Processing Email {idx}**")
                
                # Extract adjectives and CS-related sentences
                adjectives = extract_adjectives(email)
                cs_sentences = extract_cs_related_sentences(email, cs_keywords)

                # Combine features for LLM input
                combined_data = (
                    f"Adjectives found: {', '.join(adjectives)}. "
                    f"Computer science-related sentences: {' '.join(cs_sentences)}."
                )

                # Generate a structured analysis
                prompt_template = ChatPromptTemplate.from_messages(
                    [
                        (
                            "system",
                            "You are a detailed and structured feedback assistant for software engineering interns. "
                            "Based on the provided data, summarize the feedback into three sections: Strengths, Weaknesses, and Recommendations. "
                            "Ensure you use different words to paraphrase the content of the email while maintaining its meaning."
                        ),
                        ("human", "{input}"),
                    ]
                )
                
                input_data = {
                    "input": combined_data,
                }
                
                chain = prompt_template | llm
                response = chain.invoke(input_data)

                response_content = getattr(response, "content", str(response))
                structured_summary = response_content.strip()

                # Display the structured output
                st.write(f"**Feedback for Email {idx}:**")
                st.markdown(structured_summary)

                # Option to download the summarized feedback
                if st.button(f"Download Summary for Email {idx}"):
                    doc = Document()
                    doc.add_heading(f'Feedback Summary for Email {idx}', level=1)
                    doc.add_paragraph(structured_summary)
                    doc.add_paragraph(f"Adjectives: {', '.join(adjectives)}")
                    doc.add_paragraph(f"CS-related sentences: {' '.join(cs_sentences)}")
                    
                    buf = io.BytesIO()
                    doc.save(buf)
                    buf.seek(0)
                    st.download_button(
                        label="Download Summary",
                        data=buf,
                        file_name=f"email_summary_{idx}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
    else:
        st.error("No email content found in the uploaded file.")
